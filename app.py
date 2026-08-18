"""
AI Chatbot pro firmy — RAG chatbot nad nahranými dokumenty (Streamlit + LangChain + FAISS).

Odpovědi generuje LLM výhradně na základě obsahu nahraných dokumentů (PDF, TXT, DOCX, MD).
Podporovaní poskytovatelé (zobrazí se jen ti, pro které je v secrets vyplněný API klíč):
  - Google Gemini (free)
  - Groq (free, OpenAI-kompatibilní API)
  - Anthropic Claude
  - OpenAI
"""

import os
import shutil
import sys
from pathlib import Path

# Některá cloudová prostředí (např. Streamlit Community Cloud) nemají nastavenou
# UTF-8 locale, takže výchozí kódování stdout/stderr je ASCII — jakýkoliv pokus
# knihovny (httpx/openai/logging) zapsat český text pak spadne na
# UnicodeEncodeError. Vynutíme UTF-8 hned na startu, nezávisle na locale hostu.
os.environ.setdefault("PYTHONUTF8", "1")
os.environ.setdefault("PYTHONIOENCODING", "utf-8")
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(encoding="utf-8", errors="replace")

import anthropic
import google.generativeai as genai
import openai
import streamlit as st
from docx import Document as DocxReader
from google.api_core import exceptions as google_exceptions
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

# --- Konfigurace poskytovatelů ---
# Pořadí = pořadí v sidebaru. Free možnosti (Gemini, Groq) jsou první.

PROVIDERS = {
    "gemini": {
        "label": "🟢 Google Gemini (zdarma)",
        "secret_key": "GEMINI_API_KEY",
        "models": {
            # Pozn.: Gemini modely mají pravidelný deprecation cyklus (řádově měsíce) —
            # pokud model přestane fungovat, aktuální seznam je na ai.google.dev/gemini-api/docs/models.
            "Gemini 3.7 Flash (rychlý, doporučeno)": "gemini-3.7-flash",
            "Gemini 3.1 Pro (nejkvalitnější)": "gemini-3.1-pro-preview",
        },
    },
    "groq": {
        "label": "🟢 Groq (zdarma, velmi rychlý)",
        "secret_key": "GROQ_API_KEY",
        "models": {
            # Pozn.: Groq modely mají pravidelný deprecation cyklus (řádově měsíce) —
            # pokud model přestane fungovat, aktuální seznam je na console.groq.com/docs/models.
            "GPT-OSS 120B (doporučeno)": "openai/gpt-oss-120b",
            "GPT-OSS 20B (nejrychlejší)": "openai/gpt-oss-20b",
        },
    },
    "anthropic": {
        "label": "Anthropic Claude",
        "secret_key": "ANTHROPIC_API_KEY",
        "models": {
            "Claude Sonnet 5 (doporučeno)": "claude-sonnet-5",
            "Claude Opus 5 (nejvyšší kvalita)": "claude-opus-5",
            "Claude Haiku 4.5 (nejrychlejší)": "claude-haiku-4-5",
        },
    },
    "openai": {
        "label": "OpenAI",
        "secret_key": "OPENAI_API_KEY",
        "models": {
            "GPT-4o": "gpt-4o",
            "GPT-4o mini": "gpt-4o-mini",
        },
    },
}

GROQ_BASE_URL = "https://api.groq.com/openai/v1"

DEFAULT_SYSTEM_PROMPT = (
    "Jsi profesionální AI asistent pro malé firmy. Odpovídáš pouze na základě "
    "poskytnutých dokumentů. Buď stručný, přátelský a profesionální. Pokud informace "
    "nemáš, otevřeně to řekni a nabídni kontakt na člověka."
)

VECTORSTORE_DIR = Path("vectorstore_data")
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
RETRIEVAL_K = 4
MAX_FILE_SIZE_MB = 20
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

# Modely s adaptivním "thinking" zapnutým ve výchozím stavu — pro použití
# nastavení teploty je nutné thinking explicitně vypnout.
ANTHROPIC_THINKING_MODELS = {"claude-opus-5", "claude-sonnet-5"}


# --- Načtení API klíčů ---

def get_secret(key: str):
    """Bezpečně načte hodnotu z st.secrets, i pokud soubor secrets.toml neexistuje."""
    try:
        return st.secrets.get(key)
    except Exception:
        return None


API_KEYS = {key: get_secret(cfg["secret_key"]) for key, cfg in PROVIDERS.items()}
AVAILABLE_PROVIDERS = [key for key in PROVIDERS if API_KEYS[key]]


# --- Zpracování dokumentů ---

def load_pdf(file, filename: str) -> list[Document]:
    reader = PdfReader(file)
    docs = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(page_content=text, metadata={"source": filename, "page": i}))
    return docs


def load_docx(file, filename: str) -> list[Document]:
    doc = DocxReader(file)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": filename})] if text.strip() else []


def load_text(file, filename: str) -> list[Document]:
    raw = file.read()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("utf-8", errors="ignore")
    return [Document(page_content=text, metadata={"source": filename})] if text.strip() else []


def process_uploaded_files(uploaded_files) -> tuple[list[Document], list[str]]:
    """Načte nahrané soubory do LangChain Document objektů. Vrací (dokumenty, chybové hlášky)."""
    all_docs: list[Document] = []
    errors: list[str] = []

    for f in uploaded_files:
        if f.size > MAX_FILE_SIZE_BYTES:
            errors.append(f"„{f.name}“: soubor je příliš velký (max. {MAX_FILE_SIZE_MB} MB).")
            continue

        ext = Path(f.name).suffix.lower()
        try:
            if ext == ".pdf":
                docs = load_pdf(f, f.name)
            elif ext == ".docx":
                docs = load_docx(f, f.name)
            elif ext in (".txt", ".md"):
                docs = load_text(f, f.name)
            else:
                errors.append(f"„{f.name}“: nepodporovaný formát souboru.")
                continue

            if not docs:
                errors.append(f"„{f.name}“: nepodařilo se z dokumentu extrahovat žádný text.")
                continue

            all_docs.extend(docs)
        except Exception as e:
            errors.append(f"„{f.name}“: chyba při zpracování souboru ({e}).")

    return all_docs, errors


@st.cache_resource(show_spinner=False)
def get_embeddings() -> HuggingFaceEmbeddings:
    """Lokální embedding model (běží bez API klíče, výsledek se cachuje mezi requesty)."""
    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)


def load_vectorstore_from_disk(embeddings: HuggingFaceEmbeddings):
    """Načte dříve uloženou znalostní bázi z disku, pokud existuje."""
    if VECTORSTORE_DIR.exists() and any(VECTORSTORE_DIR.iterdir()):
        try:
            return FAISS.load_local(
                str(VECTORSTORE_DIR), embeddings, allow_dangerous_deserialization=True
            )
        except Exception:
            return None
    return None


def index_documents(uploaded_files) -> tuple[int, list[str]]:
    """Zpracuje a naindexuje nahrané soubory do znalostní báze (FAISS). Vrací (počet úseků, chyby)."""
    docs, errors = process_uploaded_files(uploaded_files)
    if not docs:
        return 0, errors

    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = splitter.split_documents(docs)

    embeddings = get_embeddings()
    if st.session_state.vectorstore is None:
        st.session_state.vectorstore = FAISS.from_documents(chunks, embeddings)
    else:
        st.session_state.vectorstore.add_documents(chunks)

    VECTORSTORE_DIR.mkdir(exist_ok=True)
    st.session_state.vectorstore.save_local(str(VECTORSTORE_DIR))

    return len(chunks), errors


def delete_knowledge_base():
    st.session_state.vectorstore = None
    if VECTORSTORE_DIR.exists():
        shutil.rmtree(VECTORSTORE_DIR)


# --- Vyhledávání relevantního kontextu ---

def retrieve_context(query: str) -> tuple[str, list[dict]]:
    vectorstore = st.session_state.vectorstore
    if vectorstore is None:
        return "", []

    results = vectorstore.similarity_search(query, k=RETRIEVAL_K)
    context_parts = []
    sources = []
    for doc in results:
        source = doc.metadata.get("source", "neznámý dokument")
        page = doc.metadata.get("page")
        label = f"{source}" + (f" (strana {page})" if page else "")
        context_parts.append(f"[{label}]\n{doc.page_content}")
        snippet = doc.page_content[:300].strip()
        sources.append({"label": label, "snippet": snippet + ("…" if len(doc.page_content) > 300 else "")})

    return "\n\n---\n\n".join(context_parts), sources


def build_system_prompt(base_prompt: str, context: str) -> str:
    if context.strip():
        return (
            f"{base_prompt}\n\n"
            "K dispozici máš následující výňatky z nahraných dokumentů. Odpovídej VÝHRADNĚ na "
            "základě tohoto kontextu. Pokud odpověď v kontextu nenajdeš, otevřeně řekni: "
            "„Nemám tuto informaci v dokumentech“ a nabídni kontakt na člověka. Neuváděj žádné "
            "informace, které nejsou podložené kontextem. Nepiš do odpovědi žádné interní či "
            "systémové značky (např. <thinking>).\n\n"
            f"KONTEXT Z DOKUMENTŮ:\n{context}"
        )
    return (
        f"{base_prompt}\n\n"
        "V knowledge base zatím nejsou žádné nahrané dokumenty. Informuj o tom uživatele a "
        "požádej ho, aby nejprve nahrál soubory přes postranní panel."
    )


# --- Volání LLM ---

def call_claude(model_id: str, system_prompt: str, history: list[dict], temperature: float, max_tokens: int) -> str:
    client = anthropic.Anthropic(api_key=API_KEYS["anthropic"])
    kwargs = dict(
        model=model_id,
        max_tokens=max_tokens,
        system=system_prompt,
        messages=history,
        temperature=temperature,
    )
    # Opus 5 / Sonnet 5 mají ve výchozím stavu zapnuté adaptivní "thinking", které
    # není kompatibilní s parametrem temperature — pro tento use-case ho vypínáme.
    if model_id in ANTHROPIC_THINKING_MODELS:
        kwargs["thinking"] = {"type": "disabled"}

    response = client.messages.create(**kwargs)

    if response.stop_reason == "refusal":
        detail = response.stop_details.explanation if response.stop_details else ""
        return f"Odpověď nelze z bezpečnostních důvodů vygenerovat. {detail}".strip()

    text_parts = [block.text for block in response.content if block.type == "text"]
    return "\n".join(text_parts).strip() or "Omlouvám se, nepodařilo se vygenerovat odpověď."


def call_openai_compatible(
    api_key: str,
    base_url: str | None,
    model_id: str,
    system_prompt: str,
    history: list[dict],
    temperature: float,
    max_tokens: int,
) -> str:
    """Společná implementace pro OpenAI a pro Groq (OpenAI-kompatibilní API)."""
    client = openai.OpenAI(api_key=api_key, base_url=base_url) if base_url else openai.OpenAI(api_key=api_key)
    messages = [{"role": "system", "content": system_prompt}] + history
    try:
        response = client.chat.completions.create(
            model=model_id,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
        )
    except UnicodeEncodeError as e:
        # API klíč obsahuje neascii znak — typicky proto, že v secrets zůstal
        # placeholder text (např. "tvůj-klíč") místo skutečné hodnoty klíče.
        raise ValueError(
            "API klíč obsahuje neplatné znaky. Zkontrolujte, že jste v secrets vyplnili "
            "skutečný API klíč, a ne ukázkový placeholder text."
        ) from e
    return (response.choices[0].message.content or "").strip()


def call_gemini(model_id: str, system_prompt: str, history: list[dict], temperature: float, max_tokens: int) -> str:
    genai.configure(api_key=API_KEYS["gemini"])
    model = genai.GenerativeModel(model_name=model_id, system_instruction=system_prompt)

    # Gemini používá role "user" / "model" (ne "assistant") a obsah v poli "parts".
    contents = [
        {"role": "model" if m["role"] == "assistant" else "user", "parts": [m["content"]]}
        for m in history
    ]

    response = model.generate_content(
        contents,
        generation_config=genai.types.GenerationConfig(
            temperature=temperature,
            max_output_tokens=max_tokens,
        ),
    )

    try:
        text = response.text
    except Exception:
        reason = None
        if response.prompt_feedback and response.prompt_feedback.block_reason:
            reason = response.prompt_feedback.block_reason
        return (
            "Odpověď nelze zobrazit — pravděpodobně byla zablokována bezpečnostními filtry Gemini"
            f"{f' ({reason})' if reason else ''}."
        )

    return text.strip() or "Omlouvám se, nepodařilo se vygenerovat odpověď."


def generate_answer(
    provider: str, model_id: str, system_prompt: str, history: list[dict], temperature: float, max_tokens: int
) -> tuple[str | None, str | None]:
    """Vrací (odpověď, chybová hláška) — právě jedno z toho je None."""
    try:
        if provider == "anthropic":
            return call_claude(model_id, system_prompt, history, temperature, max_tokens), None
        if provider == "openai":
            return call_openai_compatible(
                API_KEYS["openai"], None, model_id, system_prompt, history, temperature, max_tokens
            ), None
        if provider == "groq":
            return call_openai_compatible(
                API_KEYS["groq"], GROQ_BASE_URL, model_id, system_prompt, history, temperature, max_tokens
            ), None
        if provider == "gemini":
            return call_gemini(model_id, system_prompt, history, temperature, max_tokens), None
        return None, f"Neznámý poskytovatel: {provider}"

    # --- Anthropic chyby ---
    except anthropic.AuthenticationError:
        return None, "Neplatný Anthropic API klíč. Zkontrolujte hodnotu ANTHROPIC_API_KEY v secrets."
    except anthropic.PermissionDeniedError:
        return None, "Anthropic API klíč nemá oprávnění k použití tohoto modelu."
    except anthropic.NotFoundError:
        return None, "Zadaný Claude model nebyl nalezen."
    except anthropic.RateLimitError:
        return None, "Byl překročen limit požadavků na Anthropic API. Zkuste to prosím za chvíli."
    except anthropic.APIConnectionError:
        return None, "Nepodařilo se připojit k Anthropic API. Zkontrolujte internetové připojení."
    except anthropic.APIStatusError as e:
        return None, f"Chyba Anthropic API ({e.status_code}): {e.message}"

    # --- OpenAI / Groq chyby (Groq používá stejný SDK, tedy i stejné výjimky) ---
    except openai.AuthenticationError:
        key_name = "GROQ_API_KEY" if provider == "groq" else "OPENAI_API_KEY"
        return None, f"Neplatný API klíč. Zkontrolujte hodnotu {key_name} v secrets."
    except openai.RateLimitError:
        return None, "Byl překročen limit požadavků na API. Zkuste to prosím za chvíli."
    except openai.APIConnectionError:
        return None, "Nepodařilo se připojit k API. Zkontrolujte internetové připojení."
    except openai.APIStatusError as e:
        return None, f"Chyba API: {e}"
    except ValueError as e:
        return None, str(e)

    # --- Gemini chyby ---
    except google_exceptions.PermissionDenied:
        return None, "Neplatný Gemini API klíč nebo nedostatečná oprávnění. Zkontrolujte hodnotu GEMINI_API_KEY v secrets."
    except google_exceptions.Unauthenticated:
        return None, "Neplatný Gemini API klíč. Zkontrolujte hodnotu GEMINI_API_KEY v secrets."
    except google_exceptions.ResourceExhausted:
        return None, "Byl překročen limit požadavků na Gemini API (free tier). Zkuste to prosím za chvíli."
    except google_exceptions.InvalidArgument as e:
        return None, f"Neplatný požadavek na Gemini API: {e.message}"
    except (google_exceptions.DeadlineExceeded, google_exceptions.ServiceUnavailable):
        return None, "Nepodařilo se připojit ke Gemini API. Zkontrolujte internetové připojení."
    except google_exceptions.GoogleAPICallError as e:
        return None, f"Chyba Gemini API: {e.message}"

    except Exception as e:
        return None, f"Neočekávaná chyba při generování odpovědi: {e}"


# --- Streamlit UI ---

st.set_page_config(page_title="AI Chatbot pro firmy", page_icon="💬", layout="wide")

if not AVAILABLE_PROVIDERS:
    key_list = "\n".join(f'{cfg["secret_key"]} = "..."' for cfg in PROVIDERS.values())
    st.error(
        "**Chybí API klíč.**\n\n"
        "Aplikace potřebuje alespoň jeden API klíč. Doporučujeme začít s Google Gemini nebo "
        "Groq — oba mají free tarif. Přidejte klíč do souboru `.streamlit/secrets.toml` "
        "v kořeni projektu:\n\n"
        f"```toml\n{key_list}\n```\n\n"
        "Stačí vyplnit jeden řádek. Na Streamlit Cloud přidejte klíč v nastavení aplikace "
        "v sekci **Secrets**. Podrobný návod je v README.md."
    )
    st.stop()

if "messages" not in st.session_state:
    st.session_state.messages = []
if "vectorstore" not in st.session_state:
    with st.spinner("Načítám znalostní bázi..."):
        st.session_state.vectorstore = load_vectorstore_from_disk(get_embeddings())

# --- Sidebar: nastavení ---

with st.sidebar:
    st.header("⚙️ Nastavení")

    provider_labels = {key: PROVIDERS[key]["label"] for key in AVAILABLE_PROVIDERS}
    provider_choice_label = st.selectbox("Poskytovatel", list(provider_labels.values()))
    provider = next(key for key, label in provider_labels.items() if label == provider_choice_label)

    model_options = PROVIDERS[provider]["models"]
    model_label = st.selectbox("Model", list(model_options.keys()))
    model_id = model_options[model_label]

    temperature = st.slider("Teplota (kreativita)", 0.0, 1.0, 0.3, 0.05)
    max_tokens = st.slider("Max. tokenů odpovědi", 256, 4096, 1024, 128)

    system_prompt = st.text_area("Systémový prompt", value=DEFAULT_SYSTEM_PROMPT, height=170)

    st.divider()
    st.subheader("📚 Znalostní báze")

    uploaded_files = st.file_uploader(
        "Nahrát dokumenty (PDF, TXT, DOCX, MD)",
        type=["pdf", "txt", "docx", "md"],
        accept_multiple_files=True,
    )

    if st.button("📥 Zpracovat a přidat do znalostní báze", disabled=not uploaded_files, use_container_width=True):
        with st.spinner("Zpracovávám a indexuji dokumenty..."):
            n_chunks, errors = index_documents(uploaded_files)
        for err in errors:
            st.warning(err)
        if n_chunks:
            st.success(f"Přidáno {n_chunks} úseků textu z {len(uploaded_files)} souborů.")

    n_indexed = st.session_state.vectorstore.index.ntotal if st.session_state.vectorstore else 0
    st.caption(f"Aktuálně indexováno: **{n_indexed}** úseků textu")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🗑️ Smazat historii", use_container_width=True):
            st.session_state.messages = []
            st.rerun()
    with col2:
        if st.button("🗑️ Smazat KB", use_container_width=True):
            delete_knowledge_base()
            st.rerun()

# --- Hlavní chat ---

st.title("💬 AI Chatbot pro firmy")
st.caption("Odpovědi vycházejí výhradně z nahraných dokumentů.")

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        if msg.get("sources"):
            with st.expander("📄 Zdroje"):
                for s in msg["sources"]:
                    st.markdown(f"**{s['label']}**")
                    st.caption(s["snippet"])

user_input = st.chat_input("Zeptejte se na cokoliv z nahraných dokumentů...")

if user_input:
    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.markdown(user_input)

    with st.chat_message("assistant"):
        with st.spinner("Přemýšlím..."):
            context, sources = retrieve_context(user_input)
            full_system_prompt = build_system_prompt(system_prompt, context)
            history = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages]
            answer, error = generate_answer(
                provider, model_id, full_system_prompt, history, temperature, max_tokens
            )

        if error:
            st.error(error)
        else:
            st.markdown(answer)
            if sources:
                with st.expander("📄 Zdroje"):
                    for s in sources:
                        st.markdown(f"**{s['label']}**")
                        st.caption(s["snippet"])
            st.session_state.messages.append({"role": "assistant", "content": answer, "sources": sources})
